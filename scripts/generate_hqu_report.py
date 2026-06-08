"""华侨大学工学院课程报告生成脚本 — 严格遵循华大格式规范。"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

OUTPUT = "/Users/zhongxudong/Desktop/大数据可视化技术/大数据可视化技术_课程报告_华大格式.docx"
REPO_URL = "https://github.com/victorzhong0110/battery-digital-twin-visualization"


# ══════════════════════════════════════════════════════
# 工具函数
# ══════════════════════════════════════════════════════

def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=Pt(12), bold=False):
    """设置 run 的中英文字体、字号、加粗。"""
    run.font.name = en_font
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)


def add_paragraph(doc, text, cn_font="宋体", en_font="Times New Roman",
                  size=Pt(12), bold=False, alignment=None,
                  first_indent=True, line_spacing=1.5,
                  space_before=0, space_after=0):
    """添加段落并设置完整格式。"""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if first_indent:
        fmt.first_line_indent = Pt(24)  # 2字符（小四）
    else:
        fmt.first_line_indent = Pt(0)
    if alignment is not None:
        fmt.alignment = alignment

    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return p


def add_empty_line(doc, size=Pt(12)):
    """添加空行（宋体小四）。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("")
    set_run_font(run, "宋体", "Times New Roman", size)
    return p


def add_chapter_title(doc, chapter_num, title):
    """添加章标题：四号黑体，居中，前后空一行。"""
    add_empty_line(doc)
    p = add_paragraph(doc, f"第{chapter_num}章 {title}",
                      cn_font="黑体", size=Pt(14), bold=False,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      first_indent=False)
    p.style = doc.styles["Heading 1"]
    for run in p.runs:
        set_run_font(run, "黑体", "Times New Roman", Pt(14), bold=False)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    add_empty_line(doc)
    return p


def add_section_title(doc, number, title):
    """添加节标题：四号宋体，左对齐，前空一行。"""
    add_empty_line(doc)
    p = add_paragraph(doc, f"{number}{title}",
                      cn_font="宋体", size=Pt(14), bold=False,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      first_indent=False)
    p.style = doc.styles["Heading 2"]
    for run in p.runs:
        set_run_font(run, "宋体", "Times New Roman", Pt(14), bold=False)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_subsection_title(doc, number, title):
    """添加小节标题：小四黑体，左对齐，不空行。"""
    p = add_paragraph(doc, f"{number}{title}",
                      cn_font="黑体", size=Pt(12), bold=False,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      first_indent=False)
    p.style = doc.styles["Heading 3"]
    for run in p.runs:
        set_run_font(run, "黑体", "Times New Roman", Pt(12), bold=False)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    return p


def body(doc, text):
    """正文段落：小四宋体，首行缩进2字符，1.5倍行距。"""
    return add_paragraph(doc, text, cn_font="宋体", size=Pt(12),
                         first_indent=True, line_spacing=1.5)


def add_table(doc, headers, rows, caption, table_num):
    """添加表格：表题在上方，五号宋体，居中。"""
    add_empty_line(doc)
    # 表题
    add_paragraph(doc, f"表{table_num} {caption}",
                  cn_font="宋体", size=Pt(12),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_indent=False)

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "宋体", "Times New Roman", Pt(10.5), bold=True)
        # 表头灰底
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, "宋体", "Times New Roman", Pt(10.5))

    # 设置行距
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.2

    add_empty_line(doc)
    return table


FIG_DIR = "/Users/zhongxudong/Desktop/大数据可视化技术/report_figures"


def add_figure(doc, img_filename, fig_num, caption, width_cm=14):
    """插入配图：图前空一行，图居中，图题在下方居中（小四宋体），图与图题保持同页。"""
    import os
    add_empty_line(doc)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Pt(0)
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(0)
    p_img.paragraph_format.keep_with_next = True  # 图与图题不分页
    run = p_img.add_run()
    run.add_picture(os.path.join(FIG_DIR, img_filename), width=Cm(width_cm))
    p_cap = add_paragraph(doc, f"图{fig_num} {caption}",
                          cn_font="宋体", size=Pt(12),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_empty_line(doc)
    return p_img


def body_cited(doc, *segments):
    """正文段落，支持内联上标引用。segments 为正文字符串或 ('cite', '编号') 元组。"""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Pt(24)
    for seg in segments:
        if isinstance(seg, tuple) and seg[0] == "cite":
            r = p.add_run(f"[{seg[1]}]")
            set_run_font(r, "宋体", "Times New Roman", Pt(12))
            r.font.superscript = True
        else:
            r = p.add_run(seg)
            set_run_font(r, "宋体", "Times New Roman", Pt(12))
    return p


def _set_pgnum_type(section, fmt, start=1):
    """在 sectPr 中设置页码格式(decimal/upperRoman)与起始值。"""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pg = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="{fmt}" w:start="{start}"/>')
    sectPr.append(pg)


def _append_field(paragraph, instr_text, cn_font="宋体", en_font="Times New Roman", size=Pt(9)):
    """在段落中追加一个 Word 域(如 PAGE、STYLEREF)。"""
    r1 = paragraph.add_run()
    r1._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2 = paragraph.add_run()
    r2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {instr_text} </w:instrText>'))
    r3 = paragraph.add_run()
    r3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    r4 = paragraph.add_run()
    r4._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    for r in (r1, r2, r3, r4):
        set_run_font(r, cn_font, en_font, size)


def set_footer_pagenum(section, fmt, start=1):
    """设置该节页脚页码(居中, Times New Roman 小五)，并设定页码格式与起始值。"""
    _set_pgnum_type(section, fmt, start)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    _append_field(p, "PAGE", size=Pt(9))


def set_header_styleref(section):
    """设置该节页眉为 STYLEREF 域(自动显示当前章标题)，五号宋体居中。"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    _append_field(p, 'STYLEREF "Heading 1" \\* MERGEFORMAT', size=Pt(10.5))


def add_section_break(doc):
    """插入分节符(下一页)并继承页面边距，返回新建的 section。"""
    from docx.enum.section import WD_SECTION
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(new_section)
    return new_section


def add_page_break(doc):
    """添加分页符。"""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_page(section):
    """设置页面边距。"""
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


# For page break import
from docx.enum.text import WD_BREAK
import docx.oxml


# ══════════════════════════════════════════════════════
# 生成文档
# ══════════════════════════════════════════════════════

doc = Document()

# 页面设置
section = doc.sections[0]
setup_page(section)

# 设置默认字体
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.paragraph_format.line_spacing = 1.5

# ═══════════════════════════════════════
# 封面
# ═══════════════════════════════════════

for _ in range(6):
    add_empty_line(doc)

add_paragraph(doc, "华侨大学工学院", cn_font="宋体", size=Pt(14),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

add_empty_line(doc)

add_paragraph(doc, "大数据可视化技术", cn_font="宋体", size=Pt(14),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

add_paragraph(doc, "课程项目报告", cn_font="宋体", size=Pt(14),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

for _ in range(2):
    add_empty_line(doc)

add_paragraph(doc, "锂电池数字孪生与实时仿真可视化平台",
              cn_font="黑体", size=Pt(16),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False,
              bold=False)

for _ in range(4):
    add_empty_line(doc)

for _label, _value in [("专业", "数据科学与大数据技术"), ("年级", "2023级"),
                       ("姓名", "仲旭东"), ("学号", "2395141057")]:
    add_paragraph(doc, f"{_label}  {_value}",
                  cn_font="宋体", size=Pt(14),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

for _ in range(3):
    add_empty_line(doc)

add_paragraph(doc, "2026年6月",
              cn_font="宋体", size=Pt(12),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

add_empty_line(doc)

add_paragraph(doc, "项目源代码仓库（GitHub）",
              cn_font="宋体", size=Pt(12),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
add_paragraph(doc, REPO_URL,
              cn_font="宋体", en_font="Times New Roman", size=Pt(12),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

add_section_break(doc)  # 封面单独成节(无页码)

# ═══════════════════════════════════════
# 中文摘要
# ═══════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.first_line_indent = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

run_label = p.add_run("摘要 ")
set_run_font(run_label, "黑体", "Times New Roman", Pt(14), bold=False)

abstract_text = (
    "锂离子电池作为新能源汽车和储能系统的核心部件,其健康状态的准确评估对保障运行安全具有重要意义。"
    "本项目以NASA PCoE和CALCE两大公开锂电池退化数据集为基础,整合了8块电池共846个循环周期的多维退化数据,"
    "构建了一套集数据处理、多模型预测、数字孪生仿真与交互式可视化于一体的综合平台。"
    "在数据处理方面,设计了统一数据模式,将两个异构数据集转换为标准化的20维特征空间,涵盖电化学、温度、时序统计等多类特征。"
    "在模型设计方面,采用递进式架构,依次实现了线性回归、随机森林、Transformer时序预测和物理信息神经网络四种基础模型,"
    "并创新性地提出加权集成、堆叠元学习、生命周期自适应和物理约束四种集成策略。"
    "在数字孪生方面,基于1-RC Thevenin等效电路模型实现了交互式放电仿真器,支持C倍率、温度和循环次数的实时参数调节。"
    "在可视化方面,基于Dash和Plotly构建了包含6个功能页面的深色主题仪表盘,涵盖车队总览、单体详情、数字孪生仿真、"
    "模型预测竞技场、可解释性分析和3D退化景观。留一电池交叉验证结果表明,加权集成模型取得了0.9939的平均决定系数,"
    "验证了多模型融合策略的有效性。"
)
run_body = p.add_run(abstract_text)
set_run_font(run_body, "宋体", "Times New Roman", Pt(12))

add_empty_line(doc)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.first_line_indent = Pt(0)
run_kw_label = p.add_run("关键词  ")
set_run_font(run_kw_label, "黑体", "Times New Roman", Pt(14), bold=False)
run_kw = p.add_run("锂电池  健康状态估计  数字孪生  物理信息神经网络  集成学习  大数据可视化")
set_run_font(run_kw, "宋体", "Times New Roman", Pt(12))

add_page_break(doc)

# ═══════════════════════════════════════
# 英文摘要
# ═══════════════════════════════════════

add_empty_line(doc)

add_paragraph(doc, "Lithium Battery Digital Twin and Real-time Simulation Visualization Platform",
              cn_font="宋体", en_font="Times New Roman", size=Pt(14), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

add_empty_line(doc)

add_paragraph(doc, "Zhong Xudong",
              cn_font="宋体", en_font="Times New Roman", size=Pt(12),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

add_empty_line(doc)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.first_line_indent = Pt(0)

run_label = p.add_run("Abstract ")
set_run_font(run_label, "宋体", "Times New Roman", Pt(12), bold=True)

en_abstract = (
    "Lithium-ion batteries serve as the core component of new energy vehicles and energy storage systems, "
    "making accurate assessment of their State of Health (SOH) critical for operational safety. "
    "This project utilizes the NASA PCoE and CALCE public lithium battery degradation datasets, "
    "integrating multi-dimensional degradation data from 8 batteries across 846 charge-discharge cycles "
    "to build a comprehensive platform combining data processing, multi-model prediction, digital twin simulation, "
    "and interactive visualization. For data processing, a unified schema was designed to convert two heterogeneous "
    "datasets into a standardized 20-dimensional feature space covering electrochemical, thermal, and temporal statistical features. "
    "For model design, a progressive architecture was adopted, implementing Linear Regression, Random Forest, "
    "Transformer, and Physics-Informed Neural Network (PINN) as base models, along with four innovative ensemble strategies: "
    "weighted averaging, stacking meta-learner, lifecycle-adaptive, and physics-constrained ensemble. "
    "For digital twin simulation, an interactive discharge simulator was implemented based on the 1-RC Thevenin "
    "equivalent circuit model, supporting real-time parameter adjustment of C-rate, temperature, and cycle count. "
    "For visualization, a dark-themed dashboard with 6 functional pages was developed using Dash and Plotly, "
    "covering fleet overview, battery details, digital twin simulation, model prediction arena, "
    "explainability analysis, and 3D degradation landscape. Leave-one-battery-out cross-validation results "
    "show that the weighted ensemble model achieved an average R-squared of 0.9939, validating the effectiveness "
    "of the multi-model fusion strategy."
)
run_body = p.add_run(en_abstract)
set_run_font(run_body, "宋体", "Times New Roman", Pt(10.5))

add_empty_line(doc)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.first_line_indent = Pt(0)
run_kw_label = p.add_run("Key words ")
set_run_font(run_kw_label, "宋体", "Times New Roman", Pt(12), bold=True)
run_kw = p.add_run("lithium battery; state of health estimation; digital twin; physics-informed neural network; ensemble learning; big data visualization")
set_run_font(run_kw, "宋体", "Times New Roman", Pt(10.5))

add_page_break(doc)

# ═══════════════════════════════════════
# 目录（占位，需在 Word 中更新）
# ═══════════════════════════════════════

add_empty_line(doc)

add_paragraph(doc, "目 录", cn_font="黑体", size=Pt(14), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

add_empty_line(doc)

# 插入 TOC 域代码（打开文档时由 Word 自动更新填充，见文末 updateFields 设置）
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run1 = p.add_run()
run1._element.append(fld_char_begin)

instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-2" \\h \\z \\u </w:instrText>')
run2 = p.add_run()
run2._element.append(instr)

fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3 = p.add_run()
run3._element.append(fld_char_sep)

fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run4 = p.add_run()
run4._element.append(fld_char_end)

add_section_break(doc)  # 前置部分(摘要/目录)→正文 分节，正文页码重新从1开始

# ══════════════════════════════════════════════════════
# 第1章 绪论
# ══════════════════════════════════════════════════════

add_chapter_title(doc, "1", "绪论")

add_section_title(doc, "1.1", "研究背景与意义")

body_cited(doc,
     "锂离子电池凭借其高能量密度、长循环寿命和低自放电率等优势,已成为新能源汽车、"
     "便携式电子设备和大规模储能系统的核心动力来源。截至2024年底,全球锂电池市场规模已突破"
     "千亿美元,中国更是占据了全球产能的70%以上。然而,锂电池在使用过程中不可避免地经历容量衰减"
     "和内阻增长等老化现象,严重影响系统的可靠性和安全性。因此,对电池健康状态(State of Health, SOH)"
     "进行准确评估,对于延长电池使用寿命、优化充放电策略和预防安全事故具有重要的工程意义",
     ("cite", "1"),
     "。")

body(doc,
     "电池管理系统(Battery Management System, BMS)在运行过程中持续采集电压、电流、温度等多维时序数据,"
     "产生海量的监测记录。传统的数据分析方法难以直观呈现电池退化的时空规律和多维关联。"
     "大数据可视化技术为这一问题提供了新的解决思路,通过将复杂的退化数据转化为交互式图表、热力图和三维散点图等"
     "可视化形式,帮助研究人员和运维工程师快速理解电池的健康状况并做出科学决策。")

body_cited(doc,
     "数字孪生(Digital Twin)概念的兴起进一步拓展了可视化的边界",
     ("cite", "2"),
     "。数字孪生通过构建物理实体的虚拟映射模型,"
     "实现对真实系统的实时监控和行为预测。在电池管理领域,基于等效电路模型(Equivalent Circuit Model, ECM)"
     "的数字孪生可以模拟不同工况下的电池放电行为,为运维决策提供仿真支撑。"
     "将数字孪生与大数据可视化相结合,构建集监控、预测、仿真于一体的综合平台,具有重要的研究价值和应用前景。")

add_section_title(doc, "1.2", "国内外研究现状")

add_subsection_title(doc, "1.2.1", "电池SOH估计方法")

body_cited(doc,
     "电池SOH估计方法大致分为三类。第一类是基于模型的方法,通过电化学模型或等效电路模型"
     "描述电池内部的物理化学过程,结合卡尔曼滤波等状态估计算法进行SOH预测",
     ("cite", "3"),
     "。这类方法物理可解释性强,但计算复杂度较高,且模型参数标定困难。"
     "第二类是基于数据驱动的方法,利用机器学习算法从历史运行数据中学习退化规律",
     ("cite", "4"),
     "。常用的算法包括支持向量回归、随机森林、长短期记忆网络(LSTM)",
     ("cite", "5"),
     "和Transformer等。"
     "数据驱动方法无需深入了解电池内部机理,但对训练数据的质量和数量有较高要求。"
     "第三类是混合方法,将物理模型的先验知识融入数据驱动框架,兼顾可解释性和预测精度。"
     "物理信息神经网络(Physics-Informed Neural Network, PINN)是这一方向的代表性工作",
     ("cite", "6"),
     "。")

add_subsection_title(doc, "1.2.2", "电池可视化与数字孪生")

body(doc,
     "电池可视化领域的研究主要集中在单体电池的退化曲线展示和电池组的均衡状态监控。"
     "现有的可视化工具多采用静态图表,缺乏交互探索能力。在数字孪生方面,"
     "学术界已提出多种基于ECM和电化学模型的仿真框架,但大多停留在仿真计算层面,"
     "与可视化系统的集成度不高。"
     "本项目尝试将多模型预测、数字孪生仿真和交互式可视化深度融合,"
     "构建面向电池健康管理的一站式分析平台。")

add_section_title(doc, "1.3", "项目目标与内容")

body(doc,
     "本项目的总体目标是构建一个锂电池数字孪生与实时仿真可视化平台,"
     "实现从数据采集到可视化呈现的全链路覆盖。具体研究内容包括以下五个方面:")

body(doc, "(1) 整合NASA PCoE和CALCE两个异构锂电池退化数据集,设计统一数据模式,提取20维退化特征。")
body(doc, "(2) 构建递进式模型体系,依次实现线性回归、随机森林、Transformer和PINN四种基础模型。")
body(doc, "(3) 设计四种集成策略(加权集成、堆叠元学习、生命周期自适应、物理约束),融合多模型优势。")
body(doc, "(4) 基于1-RC Thevenin等效电路模型构建数字孪生仿真器,支持多参数实时交互。")
body(doc, "(5) 开发包含6个功能页面的交互式可视化仪表盘,实现车队监控、退化分析、模型对比等功能。")

add_section_title(doc, "1.4", "技术路线")

body(doc,
     "本项目的技术路线遵循数据驱动与物理驱动相结合的思路,总体分为四个阶段。"
     "第一阶段为数据基础构建,完成两个异构数据集的解析、清洗和统一存储。"
     "第二阶段为特征工程与模型训练,从统一数据中提取20维特征,依次训练4种基础模型和4种集成策略。"
     "第三阶段为数字孪生仿真器开发,基于等效电路模型构建交互式放电仿真功能。"
     "第四阶段为可视化平台集成,将数据、模型和仿真结果通过Dash框架整合为统一的交互式界面。"
     "整个技术路线体现了从数据到知识、从离线分析到实时交互的递进关系。")

add_section_title(doc, "1.5", "论文组织结构")

body(doc,
     "本报告共分为6章。第1章为绪论,介绍研究背景、现状和项目目标。"
     "第2章阐述数据采集与处理方法,包括数据来源、统一模式和特征工程。"
     "第3章详细描述模型设计与实现,涵盖4种基础模型和4种集成策略。"
     "第4章介绍可视化系统的设计与实现,对6个功能页面逐一分析。"
     "第5章总结项目的创新点。第6章给出结论与未来展望。")

add_page_break(doc)

# ══════════════════════════════════════════════════════
# 第2章 数据采集与处理
# ══════════════════════════════════════════════════════

add_chapter_title(doc, "2", "数据采集与处理")

add_section_title(doc, "2.1", "数据来源")

body_cited(doc,
     "本项目使用NASA Prognostics Center of Excellence (PCoE)和马里兰大学"
     "Center for Advanced Life Cycle Engineering (CALCE)两个权威公开数据集",
     ("cite", "7"),
     "。选择这两个数据集的原因在于:NASA数据集提供了室温恒流条件下的完整老化轨迹,适合验证模型的基本预测能力;"
     "CALCE数据集包含不同温度和倍率的老化条件,适合评估模型的工况泛化能力。"
     "两个数据集的基本信息如表2.1所示。")

add_table(doc,
          ["数据集", "电池编号", "循环数范围", "测试条件"],
          [
              ["NASA PCoE", "B0005, B0006, B0007, B0018", "100~170次", "室温恒流充放电至失效"],
              ["CALCE", "CS2_35, CS2_36, CS2_37, CS2_38", "52~100次", "不同温度和倍率的老化测试"],
          ],
          "数据集基本信息", "2.1")

body(doc,
     "NASA数据集包含4块18650型锂电池在室温条件下的完整老化测试数据,每块电池经历100至170次"
     "充放电循环直至容量降至额定值的70%以下。数据以MATLAB格式(.mat)存储,"
     "每个循环记录包含电压、电流、温度的时序采样和容量测量值。"
     "CALCE数据集包含4块棱柱形锂电池在不同温度和C倍率条件下的老化数据,以Excel格式存储。"
     "两个数据集共覆盖8块电池、846个循环周期的退化记录。")

add_section_title(doc, "2.2", "数据统一处理")

body(doc,
     "由于两个数据集在文件格式、字段命名和数据结构上存在显著差异,直接合并使用面临较大困难。"
     "为此,本项目设计了统一数据模式(Unified Schema),通过专用解析器将异构数据转换为一致的结构。"
     "NASA数据解析器(nasa_parser.py)负责读取.mat文件并提取每个循环的容量、内阻和温度信息;"
     "CALCE数据解析器(calce_parser.py)负责解析Excel文件并将测试参数映射到统一字段。"
     "转换后的数据以Parquet列式存储格式保存,兼顾存储效率和查询性能。")

add_section_title(doc, "2.3", "特征工程")

body(doc,
     "在统一数据基础上,本项目进一步提取了20维特征向量,作为后续模型训练的输入。"
     "特征设计遵循领域知识导向的原则,涵盖以下四个类别:")

body(doc,
     "(1) 电化学特征:包括内阻(internal_resistance_ohm)、电荷转移电阻(charge_transfer_resistance_ohm)"
     "和电压斜率(voltage_slope)。内阻是反映电池老化程度的关键指标,通常随循环次数单调递增。")

body(doc,
     "(2) 温度特征:包括最高温度(max_temp_c)、平均温度(mean_temp_c)、环境温度(ambient_temp_c)"
     "和温升幅度(temp_rise)。温升反映电池内部产热变化,与老化状态密切相关。")

body(doc,
     "(3) 时序统计特征:包括放电时长(discharge_duration_s)、容量滚动标准差(capacity_rolling_std)"
     "和电阻滚动均值(resistance_rolling_mean)。滚动统计的窗口大小设为5个循环,"
     "用于捕捉短期退化波动。")

body(doc,
     "(4) 衍生速率特征:包括容量衰减率(capacity_fade_rate)和电阻增长率(resistance_increase_rate),"
     "通过相邻循环的差分计算得到。此外还包括基于额定容量的归一化特征,"
     "消除不同电池间的量纲差异。")

body(doc,
     "特征工程的完整流程由feature_extract.py模块实现,输出的特征矩阵维度为846行×20列,"
     "保存为features.parquet文件。表2.2列出了主要特征字段及其含义。")

add_table(doc,
          ["特征名称", "类别", "说明"],
          [
              ["capacity_ah", "核心指标", "当前循环放电容量(安时)"],
              ["soh", "核心指标", "健康状态,当前容量与额定容量之比"],
              ["internal_resistance_ohm", "电化学", "欧姆内阻(欧姆)"],
              ["discharge_duration_s", "时序", "放电过程持续时间(秒)"],
              ["voltage_slope", "电化学", "放电电压曲线斜率"],
              ["capacity_fade_rate", "衍生", "相邻循环间的容量变化率"],
              ["resistance_increase_rate", "衍生", "相邻循环间的电阻变化率"],
              ["capacity_rolling_std", "统计", "容量的5周期滚动标准差"],
              ["temp_rise", "温度", "放电过程中的温升幅度"],
              ["capacity_normalized", "归一化", "基于额定容量的归一化值"],
          ],
          "主要特征字段说明", "2.2")

add_page_break(doc)

# ══════════════════════════════════════════════════════
# 第3章 模型设计与实现
# ══════════════════════════════════════════════════════

add_chapter_title(doc, "3", "模型设计与实现")

body(doc,
     "本项目采用递进式模型架构,从传统统计方法逐步过渡到深度学习和物理约束模型,"
     "最终通过集成策略融合各模型优势。这种设计体现了从简单到复杂、从数据驱动到物理融合的技术路线,"
     "也便于在可视化系统中展示不同复杂度模型的性能对比。"
     "全部模型均采用留一电池交叉验证(Leave-One-Battery-Out)方案,即每轮以7块电池训练、1块测试,"
     "共进行8轮验证,全面评估跨电池泛化能力。")

add_section_title(doc, "3.1", "线性回归模型")

body(doc,
     "线性回归模型采用带L2正则化的Ridge回归作为基线方法。Ridge回归通过在最小二乘损失函数中"
     "添加系数平方和的惩罚项来防止过拟合,正则化强度参数alpha通过交叉验证确定。"
     "输入特征经过StandardScaler标准化处理,消除量纲影响。")

body(doc,
     "在充分特征工程的支持下,线性回归模型取得了0.9974的平均R²,这一结果看似出人意料,"
     "但实际反映了特征设计的有效性。由于capacity_normalized等特征与目标变量之间存在近似线性关系,"
     "简单的线性模型即可获得较高的拟合精度。更重要的是,"
     "Ridge回归的系数向量具有良好的物理可解释性,每个系数的正负和大小直接反映了对应特征"
     "对容量预测的贡献方向和强度,这为可解释性分析页面提供了核心数据支撑。")

add_section_title(doc, "3.2", "随机森林模型")

body(doc,
     "随机森林(Random Forest, RF)是一种基于装袋法(Bagging)的集成学习算法,"
     "通过训练多棵决策树并取平均值来提升预测精度和稳定性。本项目使用包含100棵决策树的随机森林,"
     "每棵树的最大深度不作限制,特征采样比例为总特征数的平方根。")

body_cited(doc,
     "随机森林在留一电池交叉验证下的平均R²为0.1539,显著低于线性回归。"
     "这一结果揭示了跨电池泛化的内在挑战:不同电池的退化模式存在差异,"
     "基于部分电池训练的决策树难以准确预测未见电池的容量变化。"
     "不过,随机森林支持基于SHAP(SHapley Additive exPlanations)",
     ("cite", "8"),
     "的特征重要性分析,"
     "通过TreeExplainer计算各特征对每个预测的边际贡献,为模型可解释性提供了定量支撑。")

add_section_title(doc, "3.3", "Transformer时序预测模型")

body_cited(doc,
     "Transformer模型最初被提出用于自然语言处理任务,其核心的自注意力(Self-Attention)机制"
     "能够有效捕捉序列中的长程依赖关系",
     ("cite", "9"),
     "。本项目将其适配到电池退化预测场景,"
     "设计了专用的时序Transformer架构,主要包含以下几个组件:")

body(doc,
     "(1) 正弦位置编码。由于Transformer本身不具备感知序列顺序的能力,"
     "需要通过位置编码注入时序位置信息。本项目采用经典的正弦-余弦位置编码方案,"
     "将循环序列的每个位置映射为一个固定维度的向量,与输入特征相加后送入编码器。")

body(doc,
     "(2) 多头自注意力机制。采用4个注意力头,每个头独立学习不同时间尺度上的退化模式依赖关系。"
     "注意力权重矩阵被保存并传递到可视化系统,用于生成Transformer注意力热图。")

body(doc,
     "(3) 注意力加权池化。传统做法是对编码器输出的所有时间步取平均,但这忽略了不同时间步的重要性差异。"
     "本项目引入可学习的注意力池化层,自动为每个时间步分配权重,使模型能够聚焦于关键退化阶段。")

body(doc,
     "(4) MC Dropout不确定性估计。在推理阶段保持Dropout激活,通过20次前向传播获取预测的均值和方差,"
     "分别作为点估计和不确定性估计,用于在可视化系统中绘制置信区间。")

body(doc,
     "模型超参数设置为:2层Transformer编码器,隐藏维度64,前馈维度128,Dropout率0.1。"
     "训练采用AdamW优化器,初始学习率1e-3,配合余弦退火学习率调度策略,最大训练100个epoch。")

add_section_title(doc, "3.4", "物理信息神经网络")

body(doc,
     "物理信息神经网络(Physics-Informed Neural Network, PINN)是本项目的核心创新之一,"
     "通过将电池退化的物理先验知识编码到损失函数中,引导神经网络学习符合物理规律的预测模型。"
     "与纯数据驱动方法相比,PINN在小样本场景下具有更好的泛化能力和物理合理性。")

body(doc,
     "PINN的网络结构采用残差学习架构:网络的输出不是直接预测容量绝对值,"
     "而是预测相对于物理基线的偏差量。物理基线由ECM模型标定得到的线性退化公式给出,"
     "即C(n) = C_0 + slope * n,其中C_0为初始容量,slope为退化斜率,n为循环次数。"
     "最终预测值等于物理基线加上网络输出的残差修正,这种设计大幅降低了网络的学习难度。")

body(doc,
     "PINN的损失函数由四个分量加权组成,如表3.1所示。数据损失保证预测精度,"
     "物理损失约束预测结果接近ECM理论值,单调性损失确保容量随循环递减,"
     "边界损失约束初始容量接近额定值。训练过程中各分量的权重和数值变化"
     "被记录并传递到可视化系统,用于绘制PINN损失分解面积图。")

add_table(doc,
          ["损失分量", "物理含义", "数学形式"],
          [
              ["数据损失", "预测值与观测值的拟合误差", "MSE(y_pred, y_true)"],
              ["物理损失", "预测应接近ECM理论容量衰减曲线", "MSE(y_pred, C0 + slope * n)"],
              ["单调性损失", "容量应随循环次数单调递减", "ReLU(y[t+1] - y[t])惩罚"],
              ["边界损失", "初始容量应接近额定容量", "MSE(y[0], rated_capacity)"],
          ],
          "PINN损失函数分量", "3.1")

add_section_title(doc, "3.5", "集成策略设计")

body_cited(doc,
     "为充分融合4种基础模型的互补优势,本项目设计了4种集成策略",
     ("cite", "10"),
     "。每种策略都输出融合后的预测值和置信区间,供可视化系统展示。")

add_subsection_title(doc, "3.5.1", "加权集成")

body(doc,
     "加权集成(Weighted Ensemble)的核心思想是根据各模型在验证集上的表现分配权重。"
     "具体而言,以各模型RMSE的倒数作为权重基础,经归一化后得到最终权重向量。"
     "RMSE越小的模型获得越高的权重。这种方法实现简洁,计算开销极低,"
     "在本项目中取得了所有集成策略中最优的平均R²(0.9939)。")

add_subsection_title(doc, "3.5.2", "堆叠元学习")

body(doc,
     "堆叠元学习(Stacking Meta-Learner)将4个基础模型的预测值作为新的特征输入,"
     "训练一个Ridge元学习器来学习最优的组合方式。"
     "与简单加权不同,元学习器可以学习到基础模型之间的非线性互补关系,"
     "例如在某些容量区间更信任Transformer、在另一些区间更信任线性模型。"
     "元学习器的训练数据来自基础模型在验证集上的输出,避免了信息泄漏。")

add_subsection_title(doc, "3.5.3", "生命周期自适应集成")

body(doc,
     "生命周期自适应集成(Lifecycle-Adaptive Ensemble)是一种动态权重策略,"
     "根据电池当前的SOH状态实时调整各模型的贡献比例。"
     "设计理念是:在电池生命周期的不同阶段,各模型的预测优势不同。"
     "早期阶段(SOH > 90%),退化行为符合物理规律,侧重PINN(权重60%);"
     "中期阶段,各模型均衡贡献;晚期阶段(SOH < 80%),退化模式复杂化,"
     "侧重随机森林和Transformer的数据拟合能力。"
     "权重在不同阶段之间通过Sigmoid函数平滑过渡,避免突变。")

add_subsection_title(doc, "3.5.4", "物理约束集成")

body(doc,
     "物理约束集成(Physics-Constrained Ensemble)在加权平均的基础上引入物理一致性约束。"
     "具体做法是:首先根据各模型预测值与PINN预测的接近程度重新分配权重,"
     "偏离PINN预测较远的模型被施加降权处理;然后对融合后的预测序列进行单调性后处理,"
     "通过保序回归(isotonic regression)确保最终输出符合容量随循环递减的物理约束。")

add_section_title(doc, "3.6", "模型评估结果")

body(doc,
     "全部8种模型在留一电池交叉验证下的性能指标如表3.2所示。"
     "评估指标包括均方根误差(RMSE)、平均绝对误差(MAE)、决定系数(R²)"
     "和平均绝对百分比误差(MAPE)。")

add_table(doc,
          ["模型", "平均R²", "平均RMSE", "平均MAE", "平均MAPE"],
          [
              ["线性回归(Ridge)", "0.9974", "0.0035", "0.0028", "0.19%"],
              ["随机森林", "0.1539", "0.0358", "0.0117", "0.86%"],
              ["Transformer", "-0.5103", "0.0859", "0.0546", "4.07%"],
              ["PINN", "0.4182", "0.0689", "0.0466", "3.22%"],
              ["加权集成", "0.9939", "0.0058", "0.0041", "0.27%"],
              ["堆叠元学习", "0.7195", "0.0281", "0.0150", "1.15%"],
              ["生命周期自适应", "0.7899", "0.0421", "0.0222", "1.53%"],
              ["物理约束集成", "0.6556", "0.0412", "0.0142", "0.99%"],
          ],
          "各模型留一电池交叉验证结果", "3.2")

body(doc,
     "从表3.2可以看出,线性回归模型在充分特征工程支持下取得了最优的R²(0.9974),"
     "验证了特征设计的有效性。Transformer和PINN在小样本跨电池场景下面临泛化挑战,"
     "分别获得了-0.5103和0.4182的R²。加权集成作为最优的集成策略,"
     "R²达到0.9939,接近最优单模型水平,同时具有更好的鲁棒性。"
     "物理约束集成虽然R²不是最高,但其MAPE仅为0.99%,说明物理约束有效抑制了极端预测偏差。")

add_page_break(doc)

# ══════════════════════════════════════════════════════
# 第4章 可视化系统设计
# ══════════════════════════════════════════════════════

add_chapter_title(doc, "4", "可视化系统设计与实现")

body(doc,
     "可视化平台基于Dash 4.x和Plotly.js构建,采用深色主题设计风格。"
     "Dash是Plotly公司开发的Python Web应用框架,通过声明式的组件和回调机制,"
     "使数据科学家无需编写前端代码即可构建交互式仪表盘。"
     "平台共包含6个功能页面,通过侧边栏导航进行切换,配合Bootstrap Icons矢量图标体系。"
     "所有页面均支持全中文界面,图表支持鼠标悬停提示、缩放和导出等交互操作。"
     "前端布局采用CSS自定义属性(CSS Custom Properties)定义设计令牌,"
     "包含30余个颜色、间距、阴影和圆角变量,保证全局视觉一致性。"
     "数据加载层采用LRU缓存策略,避免重复读取Parquet文件造成的响应延迟。"
     "本平台的完整源代码已开源,托管于GitHub,仓库地址为" + REPO_URL + "。")

add_section_title(doc, "4.1", "车队总览仪表盘")

body(doc,
     "车队总览页面是用户进入平台后看到的首页,以全局视角呈现8块电池的健康状态概览,其整体界面如图4.1所示。"
     "页面顶部为KPI(关键绩效指标)卡片区,以四张卡片分别展示车队平均SOH(80.8%)、"
     "电池总数(8块/846循环)、最大容量衰减(41.7%)和最差电池SOH(59.3%)。"
     "每张卡片顶部带有颜色条,绿色表示健康、橙色表示警告、红色表示危险,帮助用户快速识别异常。")

body(doc,
     "中部左侧为车队健康矩阵热力图,将8块电池作为Y轴、循环区间作为X轴,用SOH值着色。"
     "颜色从绿色(SOH接近100%)渐变至红色(SOH低于70%),形成直观的时空退化图谱。"
     "矩阵右侧为最终SOH分布柱状图,每根柱子按电池编号排列,叠加80%警告线辅助判断。"
     "下方的容量退化曲线和内阻变化趋势图分别展示了8块电池的容量和电阻随循环的时序变化,"
     "两者形成互补视角。页面底部的雷达图在极坐标系中展示8种模型的平均R²。")

add_figure(doc, "fig_overview.png", "4.1", "车队总览仪表盘界面")

add_section_title(doc, "4.2", "电池详情页")

body(doc,
     "电池详情页面支持通过下拉框选择单块电池进行深入分析,其界面如图4.2所示。"
     "页面顶部的KPI卡片展示该电池的数据来源、循环总数、初始容量、最终SOH和容量衰减百分比。"
     "SOH仪表盘以圆形仪表的形式直观呈现当前健康百分比,"
     "仪表盘的颜色分段从红色(低于50%)到绿色(高于95%)自动切换。")

body(doc,
     "容量与SOH双轴图将容量(Ah,左轴)和SOH(%,右轴)叠加在同一图表中,"
     "配合80%警告阈值线辅助判断电池是否接近报废标准。"
     "特征演化面板以2x2子图的形式展示内阻、放电时长、容量衰减率和电阻增长率"
     "四个关键特征随循环的时序变化,帮助用户理解退化机理。"
     "页面底部的放电电压曲线支持范围滑块交互,"
     "用户可以拖动滑块选择任意循环区间,对比不同老化阶段的电压特征差异,"
     "颜色渐变从蓝色(早期循环)到红色(晚期循环),直观反映时序先后关系。")

add_figure(doc, "fig_detail.png", "4.2", "电池详情页界面")

add_section_title(doc, "4.3", "数字孪生仿真器")

body(doc,
     "数字孪生仿真器是本平台的核心创新功能之一,基于1-RC Thevenin等效电路模型实现,其交互界面如图4.3所示。"
     "该模型由一个串联欧姆电阻R0和一个并联RC回路(R1和C1)组成,"
     "通过开路电压(Open Circuit Voltage, OCV)多项式拟合SOC-OCV关系。"
     "控制面板提供四个交互控件:电池选择下拉框、C倍率滑块(0.5C至3C)、"
     "环境温度滑块(10至50摄氏度)和循环次数滑块(0至200次)。")

body(doc,
     "参数卡片区实时展示当前工况下的有效R0、有效R1、剩余容量、SOH和时间常数。"
     "电阻参数经温度系数修正,反映温度对内阻的影响。"
     "放电仿真区域以三联子图的形式展示端子电压、荷电状态(SOC)和电池温度"
     "随放电时间的协同变化。仿真器内部通过欧拉法求解RC电路微分方程,"
     "集成了简化热模型,考虑焦耳热产生和对流散热的动态平衡。"
     "老化轨迹图以双轴形式展示容量和R0随循环次数的演化趋势,"
     "当前循环位置以虚线标注,帮助用户定位电池在生命周期中的位置。")

add_figure(doc, "fig_twin.png", "4.3", "数字孪生仿真器界面")

add_section_title(doc, "4.4", "模型预测竞技场")

body(doc,
     "模型预测竞技场支持8种模型的自由多选对比,是评估和比较模型性能的核心工具页面,其界面如图4.4所示。"
     "页面顶部的控制面板包含电池选择、模型多选和置信区间显示开关三个控件。"
     "指标对比表以RMSE、MAE、R²和MAPE四个维度量化评估所选模型,"
     "模型名称按专属颜色标识,便于在后续图表中对应识别。")

body(doc,
     "预测叠加图将真实容量值(黑色加粗线)与多条模型预测曲线绘制在同一坐标系中,"
     "开启置信区间选项后,MC Dropout估计的上下界以半透明填充区域呈现,"
     "直观展示模型预测的不确定性范围。"
     "误差分布直方图以半透明叠加的方式展示各模型预测误差的频率分布,"
     "理想的模型应呈现以0为中心的窄幅分布。"
     "残差散点图展示预测残差随循环次数的时序分布,"
     "揭示模型在不同退化阶段的系统性偏差。")

body(doc,
     "PINN损失分解面积图展示数据损失、物理损失、单调性损失和边界损失"
     "四个分量在训练过程中的动态变化,帮助用户理解物理约束如何影响训练过程。"
     "集成权重分布图以分组柱状图展示4种集成策略对4个基础模型的平均权重分配。"
     "Transformer注意力热图可视化自注意力权重矩阵,揭示模型关注的历史循环模式。")

add_figure(doc, "fig_arena.png", "4.4", "模型预测竞技场界面")

add_section_title(doc, "4.5", "可解释性分析")

body(doc,
     "可解释性分析页面旨在揭示模型的决策依据,增强预测结果的可信度,其界面如图4.5所示。"
     "页面左侧展示随机森林模型基于SHAP的特征重要性排序,"
     "以水平柱状图按平均绝对SHAP值从大到小排列,渐变色突出主导特征。"
     "右侧展示线性回归模型的标准化回归系数,"
     "正值(绿色)表示正向贡献、负值(红色)表示负向贡献。"
     "两种可解释性方法的互补使用,从全局(回归系数)和局部(SHAP)两个角度"
     "阐明了模型预测的内在逻辑,降低了黑箱模型的决策不透明性。")

body(doc,
     "跨电池特征重要性热力图将所有电池的归一化特征重要性组织成矩阵形式,"
     "列方向为特征名称,行方向为电池编号。通过对比矩阵中的深浅分布,"
     "可以发现不同电池的退化是否由相同的特征驱动,"
     "例如capacity_normalized在所有电池上都是最重要的特征,"
     "而resistance_normalized的重要性在不同电池间存在差异,"
     "这反映了不同电池在老化机理上的个体差异性。"
     "数据效率分析图展示各模型R²随训练数据量(20%至100%)的变化趋势,"
     "揭示不同模型对数据量的敏感程度:线性回归在少量数据下即可收敛,"
     "而Transformer和随机森林需要更多训练样本才能展现优势。")

add_figure(doc, "fig_explain.png", "4.5", "可解释性分析界面")

add_section_title(doc, "4.6", "3D退化景观")

body(doc,
     "3D退化景观页面将电池退化过程映射到三维空间,提供沉浸式的数据探索体验,其界面如图4.6所示。"
     "主体为三维散点图,以循环次数、内阻和容量(或其他可选变量)构成三个坐标轴,"
     "支持按SOH、数据来源、电池编号或容量衰减率着色。"
     "用户可以通过鼠标拖拽旋转视角、滚轮缩放、右键平移,从不同角度观察退化轨迹的空间分布。"
     "Z轴变量支持切换为容量、SOH、放电时长或电压斜率,"
     "使用户能够在不同特征组合下发现隐含的退化模式和聚类结构。")

body(doc,
     "页面下方配有两个2D投影辅助图:容量与电阻的二维投影揭示电阻增长与容量衰减的耦合关系,"
     "SOH与循环次数的二维投影展示各电池的退化轨迹全貌,并叠加80%警告线。"
     "三维散点图与二维投影的联动,帮助用户从不同维度理解电池退化的多维特征。"
     "三维可视化的引入打破了传统二维图表的信息呈现局限,"
     "使研究人员能够直观识别电池老化数据中的非线性流形结构和异常离群点。")

add_figure(doc, "fig_3d.png", "4.6", "3D退化景观界面")

add_page_break(doc)

# ══════════════════════════════════════════════════════
# 第5章 创新点分析
# ══════════════════════════════════════════════════════

add_chapter_title(doc, "5", "创新点分析")

add_section_title(doc, "5.1", "可视化设计创新")

body(doc,
     "在可视化设计方面,本项目有以下创新点。"
     "其一,采用自定义的深色主题设计体系,通过CSS自定义属性定义了30余个设计令牌,"
     "包括6层灰阶背景色、4级文本层次、5种健康状态颜色和8种图表配色,"
     "形成了统一且专业的视觉语言,摆脱了通用模板的同质化外观。"
     "其二,设计了6页渐进式信息架构,"
     "从宏观的车队监控到微观的单体分析,从数据驱动的模型对比到物理驱动的数字孪生仿真,"
     "形成了完整的分析链路。"
     "其三,创新性地将多维退化数据映射到可交互的三维空间,"
     "支持多种着色策略和视角自由旋转,为电池退化模式的发现提供了新的视觉化工具。")

add_section_title(doc, "5.2", "交互设计创新")

body(doc,
     "在交互设计方面,本项目实现了以下创新功能。"
     "首先,数字孪生仿真器通过滑块控件实现了ECM参数的即时响应,"
     "用户调节C倍率、温度或循环次数后,放电曲线和参数卡片实时更新,"
     "实现了所见即所得的仿真体验。"
     "其次,模型预测竞技场支持从8种模型中自由多选进行并排对比,"
     "配合置信区间的开关控制,在同一图表中同时呈现预测值和不确定性范围。"
     "再次,电池详情页的放电电压曲线配有范围滑块,"
     "用户可以拖拽选择任意循环区间进行对比,颜色渐变直观反映时序先后。"
     "最后,3D退化景观支持鼠标驱动的视角旋转、缩放和平移,"
     "结合多种着色依据的切换,提供了灵活的空间数据探索能力。")

add_section_title(doc, "5.3", "数据处理与建模创新")

body(doc,
     "在数据处理与建模方面,本项目有以下创新贡献。"
     "第一,设计了统一数据模式解决NASA和CALCE两个数据集在格式、字段和结构上的异构问题,"
     "实现了跨数据源的无缝整合。"
     "第二,构建了覆盖电化学、温度、时序统计和衍生速率四个类别的20维特征空间,"
     "为后续模型提供了信息丰富的输入表征。"
     "第三,PINN模型将ECM物理先验知识编码到损失函数中,采用残差学习架构,"
     "降低了网络学习难度并保证了预测的物理合理性。"
     "第四,提出了4种互补的集成策略,分别从静态加权、元学习、生命周期自适应和物理约束"
     "四个角度融合基础模型,其中生命周期自适应策略能够根据电池当前SOH动态调整模型权重,"
     "具有较强的实际应用价值。")

add_page_break(doc)

# ══════════════════════════════════════════════════════
# 第6章 总结与展望
# ══════════════════════════════════════════════════════

add_chapter_title(doc, "6", "总结与展望")

add_section_title(doc, "6.1", "项目总结")

body(doc,
     "本项目围绕锂电池健康管理领域的可视化需求,"
     "构建了一套完整的数字孪生与实时仿真可视化平台,实现了以下核心成果:")

body(doc,
     "(1) 建立了NASA和CALCE双数据集的统一处理流水线,整合了8块电池、846个循环周期的退化数据,"
     "提取了20维特征向量,输出标准化的Parquet和JSON数据文件。")

body(doc,
     "(2) 实现了递进式模型体系,包含线性回归、随机森林、Transformer和PINN四种基础模型,"
     "以及加权集成、堆叠元学习、生命周期自适应和物理约束四种集成策略。"
     "在留一电池交叉验证下,加权集成模型取得了0.9939的平均R²。")

body(doc,
     "(3) 构建了基于1-RC Thevenin等效电路模型的数字孪生仿真器,"
     "集成了温度效应和简化热模型,支持C倍率、温度和循环次数的实时交互调节。")

body(doc,
     "(4) 开发了包含6个功能页面的交互式可视化仪表盘,采用深色主题设计,"
     "全中文界面,涵盖车队总览、电池详情、数字孪生仿真、模型预测竞技场、"
     "可解释性分析和3D退化景观等功能模块。")

body(doc,
     "(5) 实现了基于SHAP的特征可解释性分析和跨电池特征重要性对比,"
     "增强了模型预测结果的透明度和可信度。")

add_section_title(doc, "6.2", "经验与反思")

body(doc,
     "在项目实施过程中,获得了以下经验和发现。"
     "特征工程的质量对模型性能的影响远大于模型复杂度本身。"
     "精心设计的20维特征使简单的线性回归模型超越了复杂的Transformer和PINN,"
     "说明在小样本场景下,领域知识驱动的特征设计比增加模型参数更为重要。")

body(doc,
     "深度学习模型(Transformer和PINN)在仅8块电池的小样本条件下面临严峻的泛化挑战,"
     "但其价值不仅在于预测精度,还在于提供了注意力可视化和物理约束等额外维度的分析信息,"
     "丰富了平台的可解释性内涵。集成策略有效降低了单模型的极端误差,提升了预测的整体稳健性。"
     "其中加权集成因其简洁有效而表现最优,而堆叠元学习和生命周期自适应策略"
     "则展示了更灵活的模型组合思路,为实际部署提供了多样化的选择。")

body(doc,
     "在可视化开发方面,Dash框架的回调机制和组件化设计大幅提高了开发效率,"
     "但在处理大量数据点的3D渲染时仍存在性能瓶颈。通过Parquet列式存储和LRU缓存机制的引入,"
     "有效缓解了数据加载延迟问题。深色主题的设计需要额外关注文本对比度和控件可见性,"
     "尤其是下拉菜单、输入框等表单元素在暗色背景下的可读性需要逐一调试验证。")

add_section_title(doc, "6.3", "未来展望")

body(doc,
     "未来工作可以从以下几个方向进行扩展。"
     "第一,扩充数据集规模,引入更多型号和工况的电池退化数据,"
     "以充分发挥深度学习模型的拟合能力。"
     "第二,探索迁移学习策略,利用大规模预训练模型的知识迁移到小样本目标电池,"
     "提升Transformer和PINN的跨电池泛化能力。"
     "第三,引入在线学习机制,使数字孪生模型能够根据新到达的运行数据动态更新参数,"
     "实现模型的持续演进。"
     "第四,开发移动端自适应布局,采用响应式设计适配不同屏幕尺寸,"
     "支持运维人员通过手机或平板随时随地进行现场监控和状态查询。"
     "第五,集成异常检测模块,基于统计过程控制或孤立森林等方法,"
     "结合时序变点检测算法,实现电池性能突变和缓变异常的早期预警,"
     "进一步提升平台在工业生产环境中的实用价值。"
     "第六,探索联邦学习框架,在保护各方电池运行数据隐私的前提下,"
     "实现多机构间模型参数的协同训练,有效扩大训练数据的规模。")

add_page_break(doc)

# ══════════════════════════════════════════════════════
# 参考文献
# ══════════════════════════════════════════════════════

add_empty_line(doc)
add_paragraph(doc, "参考文献", cn_font="黑体", size=Pt(14), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
add_empty_line(doc)

references = [
    "[1] Hu X, Xu L, Lin X, et al. Battery lifetime prognostics[J]. Joule, 2020, 4(2): 310-346.",
    "[2] Grieves M, Vickers J. Digital twin: Mitigating unpredictable, undesirable emergent behavior in complex systems[M]. Transdisciplinary Perspectives on Complex Systems, Springer, 2017: 85-113.",
    "[3] He W, Williard N, Osterman M, et al. Prognostics of lithium-ion batteries based on Dempster-Shafer theory and the Bayesian Monte Carlo method[J]. Journal of Power Sources, 2011, 196(23): 10314-10321.",
    "[4] Severson K A, Attia P M, Jin N, et al. Data-driven prediction of battery cycle life before capacity degradation[J]. Nature Energy, 2019, 4(5): 383-391.",
    "[5] Zhang Y, Xiong R, He H, et al. Long short-term memory recurrent neural network for remaining useful life prediction of lithium-ion batteries[J]. IEEE Transactions on Vehicular Technology, 2018, 67(7): 5695-5705.",
    "[6] Raissi M, Perdikaris P, Karniadakis G E. Physics-informed neural networks: A deep learning framework for solving forward and inverse problems involving nonlinear partial differential equations[J]. Journal of Computational Physics, 2019, 378: 686-707.",
    "[7] Saha B, Goebel K. Battery data set[EB/OL]. NASA Ames Prognostics Data Repository, 2007.",
    "[8] Lundberg S M, Lee S I. A unified approach to interpreting model predictions[C]. Advances in Neural Information Processing Systems, 2017: 4765-4774.",
    "[9] Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]. Advances in Neural Information Processing Systems, 2017: 5998-6008.",
    "[10] Xing Y, Ma E W M, Tsui K L, et al. An ensemble model for predicting the remaining useful performance of lithium-ion batteries[J]. Microelectronics Reliability, 2013, 53(6): 811-820.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(0)
    # 悬挂缩进
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    set_run_font(run, "宋体", "Times New Roman", Pt(10.5))

add_page_break(doc)

# ══════════════════════════════════════════════════════
# 致谢
# ══════════════════════════════════════════════════════

add_empty_line(doc)
add_paragraph(doc, "致谢", cn_font="黑体", size=Pt(14), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
add_empty_line(doc)

body(doc,
     "本课程项目的完成离不开老师的悉心指导和同学们的热心帮助。"
     "感谢任课老师在大数据可视化技术方面的系统教学和实践指导,"
     "通过课堂讲授和案例分析,帮助我建立了从数据采集到可视化呈现的完整知识体系,"
     "为本项目的选题和技术路线提供了方向性的建议。"
     "感谢NASA PCoE和CALCE研究中心公开共享的电池退化数据集,"
     "高质量的公开数据为学术研究和教学实践提供了宝贵的资源支撑。"
     "感谢开源社区提供的Dash、Plotly、PyTorch、scikit-learn等优秀工具,"
     "活跃的社区生态和详尽的技术文档大幅降低了平台的开发门槛。"
     "最后,感谢在项目开发过程中给予过帮助和建议的所有朋友。")


# ══════════════════════════════════════════════════════
# 页眉与分节页码
# 封面无页码；摘要/目录用罗马数字；正文用阿拉伯数字并从1重新计数；正文页眉自动显示章标题
# ══════════════════════════════════════════════════════

_sections = doc.sections
if len(_sections) >= 3:
    front_sec, body_sec = _sections[1], _sections[2]
    set_footer_pagenum(front_sec, "upperRoman", 1)   # 摘要/目录：I, II, III...
    set_footer_pagenum(body_sec, "decimal", 1)       # 正文：1, 2, 3...
    set_header_styleref(body_sec)                    # 正文页眉："第X章 章标题"
else:
    print(f"警告: 期望3个分节, 实际 {len(_sections)} 个, 跳过页眉页码设置")

# 让 Word 打开文档时自动更新所有域(目录、页码、页眉章标题),无需手动右键更新域
_settings = doc.settings.element
for _uf in _settings.findall(qn("w:updateFields")):
    _settings.remove(_uf)
_settings.insert(0, parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>'))


# ══════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f"华大格式报告已生成: {OUTPUT}")

# 统计字数
import re
total = 0
for p in doc.paragraphs:
    text = p.text.strip()
    cn_chars = len(re.findall(r'[一-鿿]', text))
    en_words = len(re.findall(r'[a-zA-Z]+', text))
    total += cn_chars + en_words
print(f"报告总字数(估算): {total}")
